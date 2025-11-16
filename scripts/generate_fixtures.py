#!/usr/bin/env python3
"""
Generate sample fixture files for UI testing.

Creates datasets with duplicates, missing values, phone numbers, URLs, and outliers.
Generates data files (CSV/JSON/XLSX/Parquet/Feather), documents (DOCX/PDF/TXT), and images (PNG/JPG/TIFF/WEBP).

Usage:
  python scripts/generate_fixtures.py --output tests/fixtures --include-large --copy-to-input
"""

import argparse
from pathlib import Path
import random
import math

import pandas as pd
import numpy as np

# Optional libraries
try:
    import pyarrow  # noqa: F401
except Exception:
    pyarrow = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from PyPDF2 import PdfWriter
except Exception:
    PdfWriter = None

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None


def create_output_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def make_core_dataframe(rows: int = 1000) -> pd.DataFrame:
    rng = np.random.default_rng(42)
    df = pd.DataFrame({
        'id': np.arange(1, rows + 1),
        'name': [random.choice(['John Smith', 'Jon Smith', 'Alice', 'Bob', 'Charlie']) for _ in range(rows)],
        'age': rng.integers(18, 70, size=rows),
        'income': rng.normal(60000, 15000, size=rows).round(2),
        'phone': [random.choice(['+1 202-555-0156', '(202) 555-0143', '+44 20 7946 0958', '0044 20 7946 0958', '']) for _ in range(rows)],
        'url': [
            random.choice([
                'https://example.com/path?a=1', 'example.com/path', 'http://www.test.com/a', 'news.example.org', ''
            ]) for _ in range(rows)
        ],
    })
    # Introduce missing values
    for col in ['age', 'income']:
        mask = rng.random(rows) < 0.05
        df.loc[mask, col] = np.nan
    # Add duplicates
    dup_indices = rng.choice(rows, size=max(1, rows // 50), replace=False)
    df = pd.concat([df, df.iloc[dup_indices]], ignore_index=True)
    # Add outliers
    outlier_idx = rng.choice(len(df), size=max(1, rows // 100), replace=False)
    df.loc[outlier_idx, 'income'] = df['income'].mean() + 10 * df['income'].std()
    return df


def write_data_files(df: pd.DataFrame, out_dir: Path):
    # CSV
    df.to_csv(out_dir / 'people.csv', index=False)
    # JSON
    df.head(200).to_json(out_dir / 'people.json', orient='records', indent=2)
    # JSONL
    df.head(200).to_json(out_dir / 'people.jsonl', orient='records', lines=True)
    # TSV (as TXT)
    df.head(200).to_csv(out_dir / 'people.txt', index=False, sep='\t')
    # XML (simple)
    records = df.head(100).to_dict('records')
    xml = '<root>' + ''.join([
        '<record>' + ''.join([f'<{k}>{str(v)}</{k}>' for k, v in rec.items()]) + '</record>' for rec in records
    ]) + '</root>'
    (out_dir / 'people.xml').write_text(xml, encoding='utf-8')
    # XLSX
    df.head(5000).to_excel(out_dir / 'people.xlsx', index=False)
    # Parquet/Feather if available
    if pyarrow is not None:
        df.head(10000).to_parquet(out_dir / 'people.parquet')
        df.head(10000).to_feather(out_dir / 'people.feather')


def write_large_csv(out_dir: Path, rows: int = 12000):
    df_large = make_core_dataframe(rows)
    df_large.to_csv(out_dir / f'people_large_{rows}.csv', index=False)


def write_documents(out_dir: Path):
    # DOCX
    if docx is not None:
        d = docx.Document()
        d.add_heading('Sample Document', level=1)
        d.add_paragraph('This is a sample DOCX file for DocZilla testing.')
        d.add_paragraph('Phone: +1 202-555-0156, URL: example.com/pathhere')
        d.save(out_dir / 'sample.docx')
    # TXT
    (out_dir / 'sample.txt').write_text('Plain text file for DocZilla tests.', encoding='utf-8')
    # PDF (blank pages or simple)
    if PdfWriter is not None:
        writer = PdfWriter()
        # Add 3 blank pages (PyPDF2 cannot render text easily)
        for _ in range(3):
            writer.add_blank_page(width=612, height=792)
        with open(out_dir / 'sample.pdf', 'wb') as f:
            writer.write(f)


def write_images(out_dir: Path):
    if Image is None:
        return
    sizes = [(600, 400), (400, 600), (800, 800)]
    colors = [(200, 50, 50), (50, 200, 50), (50, 50, 200)]
    formats = ['PNG', 'JPEG', 'WEBP', 'TIFF']
    for i, (w, h) in enumerate(sizes, 1):
        img = Image.new('RGB', (w, h), color=colors[i % len(colors)])
        draw = ImageDraw.Draw(img)
        draw.text((10, 10), f'Image {i}', fill=(255, 255, 255))
        for fmt in formats:
            name = f'image_{i}.{"jpg" if fmt == "JPEG" else fmt.lower()}'
            img.save(out_dir / name, format=fmt)


def copy_to_input(fixtures_dir: Path, input_dir: Path):
    input_dir.mkdir(parents=True, exist_ok=True)
    for p in fixtures_dir.iterdir():
        if p.is_file():
            dest = input_dir / p.name
            dest.write_bytes(p.read_bytes())


def main():
    parser = argparse.ArgumentParser(description='Generate fixture files for DocZilla UI testing')
    parser.add_argument('--output', type=str, default='tests/fixtures', help='Output directory for fixtures')
    parser.add_argument('--include-large', action='store_true', help='Include large CSV to trigger auto-sampling')
    parser.add_argument('--copy-to-input', action='store_true', help='Copy generated fixtures into input/ directory')
    args = parser.parse_args()

    out_dir = create_output_dir(Path(args.output))

    # Create data
    df = make_core_dataframe(rows=1000)
    write_data_files(df, out_dir)

    # Large CSV
    if args.include_large:
        write_large_csv(out_dir, rows=12000)

    # Documents and images
    write_documents(out_dir)
    write_images(out_dir)

    # Copy to input if requested
    if args.copy_to_input:
        copy_to_input(out_dir, Path('input'))

    print(f'✓ Fixtures generated in {out_dir.resolve()}')


if __name__ == '__main__':
    main()
